import os
import re
from collections import Counter, defaultdict
from datetime import datetime,timedelta

import matplotlib
from nltk import data
# --- Third-party Libraries ---

matplotlib.use("Agg")  # Use non-GUI backend for plotting
import docx
import matplotlib.pyplot as plt
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
# --- Report Generation Libraries ---
from reportlab.platypus import (Image, Paragraph, SimpleDocTemplate, Spacer,
                                Table, TableStyle)
# -----import 3rd class---
from summary_llm import summarize_with_gemini
from textblob import TextBlob
from utils import SAMPLE_DATA_ENG, SAMPLE_DATA_VN

# ==============================================================================
# 2. REPORT CONTENT GENERATION
# ==============================================================================


def is_meaningful(content: str) -> bool:
    """Checks if content is substantial enough for analysis."""
    return len(content.split()) >= 4 and "?" not in content


def generate_overall_summary(transcript_data: list) -> str:
    """Generates a high-level executive summary of the entire meeting."""
    full_transcript = " ".join(
        [
            f"{entry['name']}: {entry['content']}"
            for entry in transcript_data
            if is_meaningful(entry["content"])
        ]
    )
    return summarize_with_gemini(
        full_transcript,
        "Provide a concise executive summary of this meeting transcript.If there use the Vietnamese, please write it in Vietnamese",
    )


def generate_key_takeaways(transcript_data: list) -> str:
    """Generates key takeaways or action items from the meeting."""
    full_transcript = " ".join(
        [
            f"{entry['name']}: {entry['content']}"
            for entry in transcript_data
            if is_meaningful(entry["content"])
        ]
    )
    return summarize_with_gemini(
        full_transcript,
        "List the key takeaways from this meeting. Use Vietnamese if appropriate, otherwise use English."
    )


def generate_speaker_summaries(transcript_data: list, speaker_durations: dict) -> dict:
    """Generates a summary of each speaker's contributions."""
    speaker_contributions = defaultdict(list)
    for entry in transcript_data:
        if is_meaningful(entry["content"]):
            speaker_contributions[entry["name"]].append(entry["content"])

    speaker_summaries = {}
    for speaker, points in speaker_contributions.items():
        combined_points = " ".join(points)
        summary = summarize_with_gemini(
            combined_points,
            f"Summarize the key points made by {speaker}.If there use the Vietnamese, please write it in Vietnamese",
        )
        speaker_summaries[speaker] = {
            "summary": summary,
            "duration": speaker_durations.get(speaker, 0),
        }
    return speaker_summaries

def generate_interval_summaries(transcript_data: list, interval_minutes: int) -> dict:
    """Generates summaries for specified time intervals."""
    if not transcript_data:
        return {}

    interval_summaries = {}
    # Use 'timeStamp' instead of 'start_time'
    start_time = datetime.fromisoformat(transcript_data[0]["timeStamp"].replace("Z", "+00:00"))
    interval_delta = timedelta(minutes=interval_minutes)
    current_interval_start = start_time
    
    # Also get the meeting end time from the last entry
    last_entry_time = datetime.fromisoformat(transcript_data[-1]["timeStamp"].replace("Z", "+00:00"))

    while current_interval_start <= last_entry_time:
        interval_end = current_interval_start + interval_delta
        interval_transcript = []

        for entry in transcript_data:
            # Use 'timeStamp' instead of 'start_time'
            entry_time = datetime.fromisoformat(entry["timeStamp"].replace("Z", "+00:00"))
            if current_interval_start <= entry_time < interval_end:
                interval_transcript.append(f"{entry['name']}: {entry['content']}")

        if interval_transcript:
            full_interval_text = " ".join(interval_transcript)
            interval_summary = summarize_with_gemini(
                full_interval_text,
                "Summarize the following conversation snippet. If Vietnamese is used, please write it in Vietnamese."
            )
            time_format = "%I:%M %p"
            interval_label = f"{current_interval_start.strftime(time_format)} - {interval_end.strftime(time_format)}"
            interval_summaries[interval_label] = interval_summary

        # Move to the next interval
        current_interval_start = interval_end
            
    return interval_summaries
# ==============================================================================
# 3. HELPER & UTILITY FUNCTIONS
# ==============================================================================


def format_time(timestamp: str) -> str:
    """Converts ISO format timestamp to a readable format with AM/PM."""
    dt = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
    return dt.strftime("%Y-%m-%d %I:%M %p")


def create_reports_directory():
    """Creates a directory named './reports' if it does not exist."""
    if not os.path.exists("./reports"):
        os.makedirs("./reports")
        print("Reports directory created successfully.")


def generate_sentiment_pie_chart(sentiment_summary: dict) -> str:
    """Generates a pie chart from sentiment data and saves it."""
    labels = list(sentiment_summary.keys())
    sizes = list(sentiment_summary.values())
    colors = ["#4CAF50", "#FFC107", "#F44336"]  # Green, Amber, Red

    plt.figure(figsize=(5, 5))
    plt.pie(
        sizes,
        labels=labels,
        colors=colors,
        autopct="%1.1f%%",
        startangle=140,
        wedgeprops={"edgecolor": "white"},
    )
    plt.axis("equal")

    chart_filename = "./reports/sentiment_pie_chart.png"
    plt.savefig(chart_filename)
    return chart_filename


def categorize_sentiment(polarity: float) -> str:
    """Categorizes sentiment based on polarity score."""
    if polarity > 0.2:
        return "Positive"
    elif polarity < -0.2:
        return "Negative"
    else:
        return "Neutral"


def analyze_speech(transcript_data: list) -> tuple[list, dict]:
    """Analyzes speech data to categorize sentiment for each entry."""
    analysis_results = []
    sentiment_summary = Counter({"Positive": 0, "Neutral": 0, "Negative": 0})

    for entry in transcript_data:
        blob = TextBlob(entry["content"])
        sentiment_category = categorize_sentiment(blob.sentiment.polarity)
        sentiment_summary[sentiment_category] += 1

        analysis_results.append(
            {
                "speaker": entry["name"],
                "content": entry["content"],
                "sentiment_category": sentiment_category,
            }
        )

    return analysis_results, dict(sentiment_summary)


# Register the font
pdfmetrics.registerFont(TTFont("DejaVuSans", "./fonts/DejaVuSans.ttf"))


def fix_style():
    styles = getSampleStyleSheet()
    for name in styles.byName:
        styles[name].fontName = "DejaVuSans"
    return styles


# ==============================================================================
# 4. REPORT CREATION FUNCTIONS (PDF & DOCX)
# ==============================================================================


# --- NORMAL REPORT ---
def create_normal_report_pdf(meeting_data):

    file_name = f"./reports/{meeting_data['meetingTitle']}_summary_report.pdf"
    doc = SimpleDocTemplate(file_name, pagesize=A4)
    styles = fix_style()
    elements = [
        Paragraph(f"<b>{meeting_data['meetingTitle']}</b>", styles["Title"]),
        Spacer(1, 12),
    ]

    # Details
    details = (
        f"<b>Convenor:</b> {meeting_data['convenor']}<br/>"
        f"<b>Start Time:</b> {format_time(meeting_data['meetingStartTimeStamp'])}<br/>"
        f"<b>End Time:</b> {format_time(meeting_data['meetingEndTimeStamp'])}<br/>"
        f"<b>Attendees:</b> {', '.join(meeting_data.get('attendees', ['N/A']))}"
    )
    elements.append(Paragraph(details, styles["Normal"]))
    elements.append(Spacer(1, 24))

    # Content
    elements.append(Paragraph("Executive Summary", styles["h2"]))
    elements.append(
        Paragraph(
            generate_overall_summary(meeting_data["transcriptData"]), styles["BodyText"]
        )
    )
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Key Takeaways", styles["h2"]))
    elements.append(
        Paragraph(
            generate_key_takeaways(meeting_data["transcriptData"]), styles["BodyText"]
        )
    )
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Speaker Summaries", styles["h2"]))
    speaker_summaries = generate_speaker_summaries(
        meeting_data["transcriptData"], meeting_data["speakerDuration"]
    )
    for speaker, data in speaker_summaries.items():
        elements.append(
            Paragraph(f"<b>{speaker}</b> ({data['duration']} seconds)", styles["h3"])
        )
        elements.append(Paragraph(data["summary"], styles["BodyText"]))
        elements.append(Spacer(1, 12))

    doc.build(elements)
    return file_name


def create_normal_report_docx(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_summary_report.docx"
    doc = docx.Document()
    doc.add_heading(meeting_data["meetingTitle"], level=1)

    # Details
    doc.add_paragraph(f"Convenor: {meeting_data['convenor']}")
    doc.add_paragraph(
        f"Time: {format_time(meeting_data['meetingStartTimeStamp'])} - {format_time(meeting_data['meetingEndTimeStamp'])}"
    )
    doc.add_paragraph(f"Attendees: {', '.join(meeting_data.get('attendees', ['N/A']))}")

    # Content
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(generate_overall_summary(meeting_data["transcriptData"]))

    doc.add_heading("Key Takeaways", level=2)
    doc.add_paragraph(generate_key_takeaways(meeting_data["transcriptData"]))

    doc.add_heading("Speaker Summaries", level=2)
    speaker_summaries = generate_speaker_summaries(
        meeting_data["transcriptData"], meeting_data["speakerDuration"]
    )
    for speaker, data in speaker_summaries.items():
        doc.add_heading(f"{speaker} ({data['duration']} seconds)", level=3)
        doc.add_paragraph(data["summary"])

    doc.save(file_name)
    return file_name


# --- SENTIMENT REPORT ---
def create_sentiment_report_pdf(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_sentiment_report.pdf"
    doc = SimpleDocTemplate(file_name, pagesize=A4)
    styles = fix_style()
    analysis, sentiment_summary = analyze_speech(meeting_data["transcriptData"])
    chart_filename = generate_sentiment_pie_chart(sentiment_summary)

    elements = [
        Paragraph("Sentiment Analysis Report", styles["Title"]),
        Spacer(1, 12),
        Image(chart_filename, width=4 * inch, height=4 * inch),
        Spacer(1, 12),
    ]

    for entry in analysis:
        elements.append(
            Paragraph(
                f"<b>{entry['speaker']}</b> ({entry['sentiment_category']})",
                styles["Normal"],
            )
        )
        elements.append(Paragraph(entry["content"], styles["BodyText"]))
        elements.append(Spacer(1, 12))

    doc.build(elements)
    return file_name


def create_sentiment_report_docx(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_sentiment_report.docx"
    doc = docx.Document()
    doc.add_heading("Sentiment Analysis Report", level=1)

    analysis, sentiment_summary = analyze_speech(meeting_data["transcriptData"])
    chart_filename = generate_sentiment_pie_chart(sentiment_summary)
    doc.add_picture(chart_filename, width=Inches(4))

    for entry in analysis:
        doc.add_heading(f"{entry['speaker']} ({entry['sentiment_category']})", level=3)
        doc.add_paragraph(entry["content"])

    doc.save(file_name)
    return file_name

# --- SPEAKER RANKING REPORT ---
def create_speaker_ranking_report_pdf(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_speaker_ranking_report.pdf"
    doc = SimpleDocTemplate(file_name, pagesize=A4)
    styles = fix_style()
    elements = [
        Paragraph("Speaker Ranking Report", styles["Title"]),
        Spacer(1, 24),
    ]

    speaker_summaries = generate_speaker_summaries(
        meeting_data["transcriptData"], meeting_data["speakerDuration"]
    )
    # Sort speakers by duration, descending
    sorted_speakers = sorted(
        speaker_summaries.items(), key=lambda item: item[1]["duration"], reverse=True
    )

    for i, (speaker, data) in enumerate(sorted_speakers, 1):
        elements.append(Paragraph(f"<b>{i}. {speaker}</b>", styles["h2"]))
        elements.append(Paragraph(f"Speaking Time: {data['duration']} seconds", styles["Normal"]))
        elements.append(Paragraph("<b>Contribution Summary:</b>", styles["Normal"]))
        elements.append(Paragraph(data["summary"], styles["BodyText"]))
        elements.append(Spacer(1, 12))

    doc.build(elements)
    return file_name


def create_speaker_ranking_report_docx(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_speaker_ranking_report.docx"
    doc = docx.Document()
    doc.add_heading("Speaker Ranking Report", level=1)

    speaker_summaries = generate_speaker_summaries(
        meeting_data["transcriptData"], meeting_data["speakerDuration"]
    )
    # Sort speakers by duration, descending
    sorted_speakers = sorted(
        speaker_summaries.items(), key=lambda item: item[1]["duration"], reverse=True
    )

    for i, (speaker, data) in enumerate(sorted_speakers, 1):
        doc.add_heading(f"{i}. {speaker}", level=2)
        doc.add_paragraph(f"Speaking Time: {data['duration']} seconds")
        doc.add_paragraph("Contribution Summary:")
        doc.add_paragraph(data["summary"])

    doc.save(file_name)
    return file_name


# --- INTERVAL REPORT ---
def create_interval_report_pdf(meeting_data, interval_minutes):
    file_name = f"./reports/{meeting_data['meetingTitle']}_interval_report.pdf"
    doc = SimpleDocTemplate(file_name, pagesize=A4)
    styles = fix_style()
    elements = [
        Paragraph(f"Interval Report ({interval_minutes}-Minute Intervals)", styles["Title"]),
        Spacer(1, 24),
    ]

    interval_summaries = generate_interval_summaries(
        meeting_data["transcriptData"], interval_minutes
    )

    if not interval_summaries:
        elements.append(Paragraph("No conversations to report.", styles["Normal"]))
    else:
        for interval, summary in interval_summaries.items():
            elements.append(Paragraph(f"<b>Interval: {interval}</b>", styles["h2"]))
            elements.append(Paragraph(summary, styles["BodyText"]))
            elements.append(Spacer(1, 12))

    doc.build(elements)
    return file_name


def create_interval_report_docx(meeting_data, interval_minutes):
    file_name = f"./reports/{meeting_data['meetingTitle']}_interval_report.docx"
    doc = docx.Document()
    doc.add_heading(f"Interval Report ({interval_minutes}-Minute Intervals)", level=1)

    interval_summaries = generate_interval_summaries(
        meeting_data["transcriptData"], interval_minutes
    )

    if not interval_summaries:
        doc.add_paragraph("No conversations to report.")
    else:
        for interval, summary in interval_summaries.items():
            doc.add_heading(f"Interval: {interval}", level=2)
            doc.add_paragraph(summary)

    doc.save(file_name)
    return file_name

# ... Add other report generation functions (SpeakerRanking, Interval) in a similar refactored style ...

# ==============================================================================
# 5. MAIN REPORT DISPATCHER
# ==============================================================================


def generate_reports(
    meeting_data, report_type="Normal", format_type="PDF", interval_minutes=5
):
    """
    Generates and saves a report based on specified choices.

    Args:
        meeting_data (dict): Contains meeting details and transcript data.
        report_type (str): "Normal", "SpeakerRanking", "Sentiment", "Interval".
        format_type (str): "PDF", "DOCX".
        interval_minutes (int): The interval in minutes for interval reports.
    """
    create_reports_directory()

    report_functions = {
        ("Normal", "PDF"): create_normal_report_pdf,
        ("Normal", "DOCX"): create_normal_report_docx,
        ("Sentiment", "PDF"): create_sentiment_report_pdf,
        ("Sentiment", "DOCX"): create_sentiment_report_docx,
        ("SpeakerRanking", "PDF"): create_speaker_ranking_report_pdf, 
        ("SpeakerRanking", "DOCX"): create_speaker_ranking_report_docx,  
        ("Interval", "PDF"): lambda data: create_interval_report_pdf(data, interval_minutes), 
        ("Interval", "DOCX"):lambda data:  create_interval_report_docx(data, interval_minutes),  
    }

    func = report_functions.get((report_type, format_type))

    if func:
        print(f"Generating {report_type} report in {format_type} format...")
        try:
            # Logic for interval reports could be added here if needed
            file_path = func(meeting_data)
            print(f"Successfully generated: {file_path}")
            return file_path
        except Exception as e:
            print(f"Failed to generate report. Error: {e}")
            return None
    else:
        raise ValueError(
            f"Invalid report/format combination: {report_type}/{format_type}"
        )


# ==============================================================================
# 6. EXAMPLE USAGE
# ==============================================================================

if __name__ == "__main__":
    # Sample meeting data object for testing

    # --- Generate all reports ---
    print("\n--- Starting Report Generation ---")
    generate_reports(SAMPLE_DATA_VN, report_type="Interval", format_type="PDF" )
    print("\n--- Report Generation Complete ---")
    print("Check the './reports' directory for the output files.")
